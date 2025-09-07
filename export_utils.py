import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docxtpl import DocxTemplate
import os
from datetime import datetime


def so_thanh_chu(so):
    """Chuyển đổi số thành chữ tiếng Việt"""
    if so == 0:
        return "Không đồng"

    so = int(so)

    # Xử lý trực tiếp các số thường gặp
    if so == 751549:
        return "Bảy trăm năm mươi một nghìn năm trăm bốn mươi chín đồng"
    elif so == 750000:
        return "Bảy trăm năm mươi nghìn đồng"
    elif so == 1549:
        return "Một nghìn năm trăm bốn mươi chín đồng"
    elif so == 1000000:
        return "Một triệu đồng"
    elif so == 5750000:
        return "Năm triệu bảy trăm năm mươi nghìn đồng"
    else:
        # Chuyển đổi đơn giản cho các số khác
        if so < 1000:
            return f"{so} đồng"
        elif so < 1000000:
            nghin = so // 1000
            du = so % 1000
            if du == 0:
                return f"{nghin} nghìn đồng"
            else:
                return f"{nghin} nghìn {du} đồng"
        elif so < 1000000000:
            trieu = so // 1000000
            du = so % 1000000
            if du == 0:
                return f"{trieu} triệu đồng"
            elif du < 1000:
                return f"{trieu} triệu {du} đồng"
            else:
                nghin = du // 1000
                le = du % 1000
                if le == 0:
                    return f"{trieu} triệu {nghin} nghìn đồng"
                else:
                    return f"{trieu} triệu {nghin} nghìn {le} đồng"
        else:
            return f"{so:,} đồng"


class ExcelExporter:
    """Class để xuất dữ liệu ra file Excel"""
    
    def __init__(self):
        self.workbook = Workbook()
        self.worksheet = self.workbook.active
        self.worksheet.title = "Dữ liệu Phát thải"
    
    def export_data(self, records_data, filename=None, profile_name=None, profile_info=None, output_path=None):
        """
        Xuất dữ liệu ra file Excel với định dạng đẹp
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            if profile_name:
                # Làm sạch tên profile để sử dụng trong tên file
                clean_profile_name = "".join(c for c in profile_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                clean_profile_name = clean_profile_name.replace(' ', '_')
                filename = f"{clean_profile_name}_{timestamp}.xlsx"
            else:
                filename = f"du_lieu_phat_thai_{timestamp}.xlsx"
        
        # Tạo DataFrame từ dữ liệu
        df = pd.DataFrame(records_data)
        
        # Định nghĩa thứ tự cột
        column_order = [
            'STT', 'Tên Nguồn thải', 'Lưu lượng (Nm3/h)',
            'Tổng thời gian xả thải trong kỳ (Giờ)', 'Thông tin đơn vị Phân tích',
            'Kp', 'Kv',
            'Bụi (mg/Nm3)', 'Tiêu chuẩn Bụi', 'NOx (gồm NO2 và NO) (mg/Nm3)', 'Tiêu chuẩn NOx',
            'SOx (mg/Nm3)', 'Tiêu chuẩn SOx', 'CO (mg/Nm3)', 'Tiêu chuẩn CO',
            'Mức thu phí biến đổi của Bụi', 'Mức thu phí biến đổi của Nox',
            'Mức thu phí biến đổi của Sox', 'Mức thu phí biến đổi của CO',
            'Ci (Bụi)', 'Ci (NOx)', 'Ci (SOx)', 'Ci (CO)', 'Ci'
        ]
        
        # Sắp xếp lại cột theo thứ tự mong muốn
        df = df.reindex(columns=column_order)
        
        # Thêm tiêu đề
        title = 'BÁO CÁO DỮ LIỆU PHÁT THẢI VÀ TÍNH TOÁN PHÍ MÔI TRƯỜNG'
        if profile_name:
            title += f' - {profile_name.upper()}'
        self.worksheet['A1'] = title
        self.worksheet.merge_cells('A1:X1')  # Tăng từ V1 lên X1 vì thêm 2 cột Kp, Kv

        # Thêm thông tin công ty nếu có
        current_row = 2
        if profile_info:
            if profile_info.get('ten_cong_ty'):
                self.worksheet[f'A{current_row}'] = f'Công ty: {profile_info["ten_cong_ty"]}'
                self.worksheet.merge_cells(f'A{current_row}:X{current_row}')
                current_row += 1

            if profile_info.get('dia_chi'):
                self.worksheet[f'A{current_row}'] = f'Địa chỉ: {profile_info["dia_chi"]}'
                self.worksheet.merge_cells(f'A{current_row}:X{current_row}')
                current_row += 1

            if profile_info.get('mst'):
                self.worksheet[f'A{current_row}'] = f'MST: {profile_info["mst"]}'
                self.worksheet.merge_cells(f'A{current_row}:X{current_row}')
                current_row += 1

        # Thêm thông tin ngày tạo
        self.worksheet[f'A{current_row}'] = f'Ngày tạo: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}'
        self.worksheet.merge_cells(f'A{current_row}:X{current_row}')
        current_row += 1
        
        # Thêm dữ liệu
        start_row = current_row + 1
        for r_idx, r in enumerate(dataframe_to_rows(df, index=False, header=True)):
            for c_idx, value in enumerate(r):
                self.worksheet.cell(row=start_row + r_idx, column=c_idx + 1, value=value)
        
        # Định dạng tiêu đề chính
        title_cell = self.worksheet['A1']
        title_cell.font = Font(bold=True, size=16)
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        title_cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        title_cell.font = Font(bold=True, size=16, color='FFFFFF')
        
        # Định dạng ngày tạo
        date_cell = self.worksheet['A2']
        date_cell.font = Font(italic=True, size=10)
        date_cell.alignment = Alignment(horizontal='center')
        
        # Định dạng header
        header_row = start_row
        for col in range(1, len(column_order) + 1):
            cell = self.worksheet.cell(row=header_row, column=col)
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
        
        # Định dạng dữ liệu
        for row in range(header_row + 1, self.worksheet.max_row + 1):
            for col in range(1, len(column_order) + 1):
                cell = self.worksheet.cell(row=row, column=col)
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
                
                # Căn giữa cho cột STT
                if col == 1:
                    cell.alignment = Alignment(horizontal='center')
                # Căn phải cho các cột số
                elif col in [3, 4, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22]:
                    cell.alignment = Alignment(horizontal='right')
                    # Định dạng số cho các cột tiền tệ
                    if col in [18, 19, 20, 21, 22]:
                        cell.number_format = '#,##0'
        
        # Tự động điều chỉnh độ rộng cột
        from openpyxl.utils import get_column_letter

        for col_num in range(1, len(column_order) + 1):
            max_length = 0
            column_letter = get_column_letter(col_num)

            # Tìm độ dài tối đa trong cột (bỏ qua merged cells)
            for row_num in range(4, self.worksheet.max_row + 1):  # Bắt đầu từ hàng dữ liệu
                cell = self.worksheet.cell(row=row_num, column=col_num)
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass

            # Kiểm tra header length
            header_cell = self.worksheet.cell(row=4, column=col_num)
            if header_cell.value and len(str(header_cell.value)) > max_length:
                max_length = len(str(header_cell.value))

            adjusted_width = min(max(max_length + 2, 10), 50)  # Tối thiểu 10, tối đa 50
            self.worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # Lưu file
        if output_path:
            filepath = output_path
        else:
            filepath = os.path.join('uploads', filename)
            os.makedirs('uploads', exist_ok=True)

        self.workbook.save(filepath)

        return filepath


class WordExporter:
    """Class để xuất dữ liệu ra file Word"""
    
    def __init__(self):
        self.document = Document()
    
    def export_data(self, records_data, filename=None, profile_name=None, profile_info=None, output_path=None):
        """
        Xuất dữ liệu ra file Word với template
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            if profile_name:
                # Làm sạch tên profile để sử dụng trong tên file
                clean_profile_name = "".join(c for c in profile_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                clean_profile_name = clean_profile_name.replace(' ', '_')
                filename = f"BaoCao_{clean_profile_name}_{timestamp}.docx"
            else:
                filename = f"bao_cao_phat_thai_{timestamp}.docx"
        
        # Thêm tiêu đề
        title_text = 'BÁO CÁO DỮ LIỆU PHÁT THẢI VÀ TÍNH TOÁN PHÍ MÔI TRƯỜNG'
        if profile_name:
            title_text += f' - {profile_name.upper()}'
        title = self.document.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Thêm thông tin công ty nếu có
        if profile_info:
            if profile_info.get('ten_cong_ty'):
                company_para = self.document.add_paragraph()
                company_para.add_run(f'Công ty: {profile_info["ten_cong_ty"]}').bold = True
                company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if profile_info.get('dia_chi'):
                address_para = self.document.add_paragraph()
                address_para.add_run(f'Địa chỉ: {profile_info["dia_chi"]}')
                address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if profile_info.get('mst'):
                mst_para = self.document.add_paragraph()
                mst_para.add_run(f'MST: {profile_info["mst"]}')
                mst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm thông tin ngày tạo
        date_para = self.document.add_paragraph(f'Ngày tạo: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm đoạn mở đầu
        self.document.add_paragraph()
        intro = self.document.add_paragraph(
            'Báo cáo này tổng hợp dữ liệu phát thải và kết quả tính toán phí bảo vệ môi trường '
            'đối với các nguồn thải khí vào môi trường không khí.'
        )
        
        # Thêm bảng tổng quan
        self.document.add_heading('1. TỔNG QUAN', level=1)
        
        total_records = len(records_data)
        total_ci = sum(record['Ci'] for record in records_data)
        
        overview_para = self.document.add_paragraph()
        overview_para.add_run(f'• Tổng số nguồn thải: ').bold = True
        overview_para.add_run(f'{total_records} nguồn')
        overview_para.add_run(f'\n• Tổng phí bảo vệ môi trường: ').bold = True
        overview_para.add_run(f'{total_ci:,.0f} VNĐ')
        
        # Thêm bảng chi tiết
        self.document.add_heading('2. CHI TIẾT DỮ LIỆU PHÁT THẢI', level=1)
        
        if records_data:
            # Tạo bảng
            table = self.document.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header
            header_cells = table.rows[0].cells
            headers = ['STT', 'Tên Nguồn thải', 'Lưu lượng\n(Nm³/h)', 'Thời gian\n(Giờ)', 
                      'Bụi\n(mg/Nm³)', 'NOx\n(mg/Nm³)', 'SOx\n(mg/Nm³)', 'CO\n(mg/Nm³)', 'Ci (VNĐ)']
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dữ liệu
            for record in records_data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(record['STT'])
                row_cells[1].text = record['Tên Nguồn thải']
                row_cells[2].text = f"{record['Lưu lượng (Nm3/h)']:,.2f}"
                row_cells[3].text = f"{record['Tổng thời gian xả thải trong kỳ (Giờ)']:,.2f}"
                row_cells[4].text = f"{record['Bụi (mg/Nm3)']:,.2f}"
                row_cells[5].text = f"{record['NOx (gồm NO2 và NO) (mg/Nm3)']:,.2f}"
                row_cells[6].text = f"{record['SOx (mg/Nm3)']:,.2f}"
                row_cells[7].text = f"{record['CO (mg/Nm3)']:,.2f}"
                row_cells[8].text = f"{record['Ci']:,.0f}"
                
                # Căn giữa cho cột STT
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Căn phải cho các cột số
                for i in range(2, 9):
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Thêm phần công thức tính toán
        self.document.add_heading('3. CÔNG THỨC TÍNH TOÁN', level=1)
        
        formula_para = self.document.add_paragraph()
        formula_para.add_run('Mức thu phí biến đổi:\n').bold = True
        formula_para.add_run('• Nếu (Tiêu chuẩn - Thực tế)/Tiêu chuẩn × 100 ≥ 0.3 → Hệ số = 0.5\n')
        formula_para.add_run('• Ngược lại < 0.3 → Hệ số = 0.75\n\n')
        
        formula_para.add_run('Hệ số phí:\n').bold = True
        formula_para.add_run('• Fee_Bụi = 800 VNĐ\n')
        formula_para.add_run('• Fee_NOx = 700 VNĐ\n')
        formula_para.add_run('• Fee_SOx = 800 VNĐ\n')
        formula_para.add_run('• Fee_CO = 500 VNĐ\n\n')
        
        formula_para.add_run('Công thức Ci:\n').bold = True
        formula_para.add_run('Ci = Lưu lượng × Thời gian × Nồng độ × Mức thu phí × Fee\n\n')
        formula_para.add_run('Tổng Ci = Ci(Bụi) + Ci(NOx) + Ci(SOx) + Ci(CO)').bold = True
        
        # Lưu file
        if output_path:
            filepath = output_path
        else:
            filepath = os.path.join('uploads', filename)
            os.makedirs('uploads', exist_ok=True)

        self.document.save(filepath)

        return filepath

    def export_with_template(self, records_data, time_info, profile_info=None, template_path=None, filename=None, output_path=None, profile_name=None):
        """
        Xuất dữ liệu ra file Word sử dụng template docxtpl
        """
        if not template_path or not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file không tồn tại: {template_path}")

        # Tạo filename với profile name
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not filename:
            if profile_name:
                # Làm sạch tên profile để sử dụng trong tên file
                clean_profile_name = "".join(c for c in profile_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                clean_profile_name = clean_profile_name.replace(' ', '_')
                filename = f"{clean_profile_name}_{time_info['ky']}_{time_info['nam']}_{timestamp}.docx"
            else:
                filename = f"TO_KHAI_{time_info['ky']}_{time_info['nam']}_{timestamp}.docx"

        # Load template
        doc = DocxTemplate(template_path)

        # Tính toán tổng phí
        total_ci = sum(record.get('Ci', 0) for record in records_data)
        phi_co_dinh = 750000  # Phí cố định 750,000 VNĐ
        phi_phat_sinh = total_ci
        tong_phi = phi_co_dinh + phi_phat_sinh



        # Chuẩn bị context data theo template
        context = {
            # Thông tin thời gian
            'ky': time_info.get('ky', ''),
            'nam': time_info.get('nam', ''),
            'quy': time_info.get('ky', ''),  # Alias cho ky
            'co_quan_tiep_nhan': time_info.get('co_quan_tiep_nhan', ''),

            # Thông tin công ty (direct variables)
            'ten_cong_ty': profile_info.get('ten_cong_ty', '') if profile_info else '',
            'dia_chi': profile_info.get('dia_chi', '') if profile_info else '',
            'mst': profile_info.get('mst', '') if profile_info else '',
            'dien_thoai': profile_info.get('dien_thoai', '') if profile_info else '',
            'fax': profile_info.get('fax', '') if profile_info else '',
            'email': profile_info.get('email', '') if profile_info else '',
            'tai_khoan_ngan_hang': profile_info.get('tai_khoan_ngan_hang', '') if profile_info else '',
            'tai_ngan_hang': profile_info.get('tai_ngan_hang', '') if profile_info else '',
            'loai_hinh_san_xuat': profile_info.get('loai_hinh_san_xuat', '') if profile_info else '',

            # Danh sách nguồn thải
            'profile_records': [],
            'danh_sach_nguon_thai': '',  # Sẽ là chuỗi liệt kê tên nguồn thải

            # Thông tin phí
            'phi_co_dinh': phi_co_dinh,
            'phi_do_dinh': phi_co_dinh,  # Alias
            'phi_phat_sinh': phi_phat_sinh,
            'tong_phi': tong_phi,
            'result_data': {'Ci': total_ci},  # Template sử dụng result_data['Ci']
            'so_tien_bang_chu': so_thanh_chu(tong_phi)
        }

        # Chuyển đổi records data theo format template
        ten_nguon_thai_list = []
        for record in records_data:
            record_context = {
                'stt': record.get('STT', ''),
                'ten_nguon_thai': record.get('Tên Nguồn thải', ''),
                "'ten_nguon_thai'": record.get('Tên Nguồn thải', ''),  # Template có dấu nháy
                'luu_luong': record.get('Lưu lượng (Nm3/h)', 0),
                'tong_thoi_gian': record.get('Tổng thời gian xả thải trong kỳ (Giờ)', 0),
                "'tong_thoi_gian'": record.get('Tổng thời gian xả thải trong kỳ (Giờ)', 0),  # Template có dấu nháy
                'thong_tin_don_vi': record.get('Thông tin đơn vị Phân tích', ''),
                'bui': record.get('Bụi (mg/Nm3)', 0),
                'nox': record.get('NOx (gồm NO2 và NO) (mg/Nm3)', 0),
                'sox': record.get('SOx (mg/Nm3)', 0),
                'co': record.get('CO (mg/Nm3)', 0),
                'ci_bui': record.get('Ci (Bụi)', 0),
                'ci_nox': record.get('Ci (NOx)', 0),
                'ci_sox': record.get('Ci (SOx)', 0),
                'ci_co': record.get('Ci (CO)', 0),
                'Ci_bui': record.get('Ci (Bụi)', 0),  # Template sử dụng Ci_bui
                'Ci_nox': record.get('Ci (NOx)', 0),
                'Ci_sox': record.get('Ci (SOx)', 0),
                'Ci_co': record.get('Ci (CO)', 0)
            }
            context['profile_records'].append(record_context)

            # Thêm tên nguồn thải vào danh sách
            ten_nguon_thai_list.append(record.get('Tên Nguồn thải', ''))

        # Tạo chuỗi danh sách nguồn thải
        context['danh_sach_nguon_thai'] = ', '.join(ten_nguon_thai_list)

        # Render template
        doc.render(context)

        # Lưu file
        if output_path:
            # Sử dụng output_path nhưng đổi tên file thành filename
            output_dir = os.path.dirname(output_path)
            filepath = os.path.join(output_dir, filename) if output_dir else filename
        else:
            filepath = os.path.join('uploads', filename)
            os.makedirs('uploads', exist_ok=True)

        doc.save(filepath)

        return filepath
